"""Shared file content extractor used by both API routes and the processor."""
import os
import logging

logger = logging.getLogger(__name__)


def extract_text_from_bytes(content: bytes, ext: str) -> str:
    """
    Extract plain text from raw bytes of a .docx or .pdf file.
    """
    import io
    if not content:
        return ""

    ext = ext.lower().lstrip(".")
    try:
        if ext == "pdf":
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages).strip()

        if ext in ("docx", "doc"):
            import docx
            doc = docx.Document(io.BytesIO(content))
            text_blocks = []
            
            for p in doc.paragraphs:
                if p.text.strip():
                    text_blocks.append(p.text.strip())
                    
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip().replace('\n', ' '))
                    if row_text:
                        text_blocks.append(" ".join(row_text))
                        
            return "\n".join(text_blocks).strip()

    except Exception as e:
        logger.error(f"extract_text_from_bytes failed: {e}")

    return ""
